import json
import urllib.request
import re

from docx import Document
from docx.shared import Inches


class Doc:
    def __init__(self, conf, name):
        self.c = conf
        self.d = Document()
        self.n = name

    def queryip(self, ip, port, sql):
        pattern = re.compile(r'\winstance5')
        # sub = pattern.findall("/*instance*/")
        sql = pattern.sub(ip + "%3A" + port, sql)
        link = "http://localhost:8080/api/v1/query?query=" + sql
        r = urllib.request.urlopen(link)
        res = r.read().decode("unicode_escape")
        j = json.loads(res)['data']['result'][0]['value'][1]
        return j

    def query(self, sql):
        link = "http://localhost:9090/api/v1/query?query=" + sql
        r = urllib.request.urlopen(link)
        res = r.read().decode("unicode_escape")
        j = json.loads(res)['data']['result'][0]['value'][1]
        return j

    def cover(self):
        self.d.add_heading(self.n, 0)
        mark, description, status, note = [], [], [], []
        for level in self.c:
            if level == "show":
                self.d.add_heading("一、系统资源及运行状态", level=1)
                for a in self.c[level]:
                    app = self.c[level][a]
                    self.d.add_heading(app['name'], level=2)
                    self.d.add_paragraph(app['ip'], style='Intense Quote')
                    for s in app['status']:
                        row = app['status'][s]
                        mark.append(s)
                        description.append(row['description'])
                        status.append(self.queryip(app['ip'], app['telegraf'], row['sql']))
                        note.append(row['note'])
                    self.table(mark, description, status, note)
                    mark, description, status, note = [], [], [], []
            elif level == "other":
                self.d.add_heading("二、系统资源及运行状态", level=1)
                for a in self.c[level]:
                    app = self.c[level][a]
                    self.d.add_heading(app['name'], level=2)
                    for s in app['status']:
                        row = app['status'][s]
                        mark.append(s)
                        description.append(row['description'])
                        status.append(self.query(row['sql']))
                        note.append(row['note'])
                    self.table(mark, description, status, note)
                    mark, description, status, note = [], [], [], []
        self.d.add_heading("三、运维工作说明", level=1)
        self.d.save('demo.docx')

    def table(self, mark, description, status, note):
        id = 1
        t = self.d.add_table(rows=1, cols=5)
        hdr_cells = t.rows[0].cells
        hdr_cells[0].text = '序号'
        hdr_cells[1].text = '资源标识'
        hdr_cells[2].text = '资源简述'
        hdr_cells[3].text = '运行状态'
        hdr_cells[4].text = '说明'
        for m, d, s, n in zip(mark, description, status, note):
            row_cells = t.add_row().cells
            row_cells[0].text = str(id)
            row_cells[1].text = m
            row_cells[2].text = d
            row_cells[3].text = s
            row_cells[4].text = n
            id += 1

# document = Document()
#
# document.add_heading('Document Title', 0)
#
# p = document.add_paragraph('A plain paragraph having some ')
# p.add_run('bold').bold = True
# p.add_run(' and some ')
# p.add_run('italic.').italic = True
#
# document.add_heading('Heading, level 1', level=1)
# document.add_paragraph('Intense quote', style='Intense Quote')
#
# document.add_paragraph(
#     'first item in unordered list', style='List Bullet'
# )
# document.add_paragraph(
#     'first item in ordered list', style='List Number'
# )
#
# # document.add_picture('monty-truth.png', width=Inches(1.25))
#
# records = (
#     (3, '101', 'Spam'),
#     (7, '422', 'Eggs'),
#     (4, '631', 'Spam, spam, eggs, and spam')
# )
#
# table = document.add_table(rows=1, cols=3)
# hdr_cells = table.rows[0].cells
# hdr_cells[0].text = 'Qty'
# hdr_cells[1].text = 'Id'
# hdr_cells[2].text = 'Desc'
# for qty, id, desc in records:
#     row_cells = table.add_row().cells
#     row_cells[0].text = str(qty)
#     row_cells[1].text = id
#     row_cells[2].text = desc
#
# document.add_page_break()
